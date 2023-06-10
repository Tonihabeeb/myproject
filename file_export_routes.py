from flask import Blueprint, send_file, Response
from docx import Document
from io import BytesIO
from datetime import datetime
from flask_jwt_extended import jwt_required, get_jwt_identity
from database import db
from document_generation import generate_document_content

file_export_routes = Blueprint('file_export_routes', __name__)

@file_export_routes.route('/export/<task_id>', methods=['GET'])
@jwt_required()
def export_task_data(task_id):
    user_id = get_jwt_identity()

    # Fetch task data
    task_data = db.find_one('tasks', {'task_id': task_id, 'user_id': user_id})
    
    if task_data is None:
        return Response("Task not found", status=404)

    # Extract task name
    task_name = task_data.get('name', 'Unnamed Task')

    # Fetch chat data for the task
    chat_data = list(db.find('ChatMessages', {"taskId": task_id}))

    # Generate document content using GPT-4 and the chat data
    if chat_data:
        document_content = generate_document_content(task_data, chat_data)
    else:
        document_content = f"No data found for task '{task_name}'"

    # Create a new Document and add the task name as a header
    doc = Document()
    doc.add_heading(task_name, level=1)

    # Add the generated content
    doc.add_paragraph(document_content)
    
    # Save the document to a BytesIO object
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    # Get current date and time
    now = datetime.now()

    # Format as string
    dt_string = now.strftime("%Y-%m-%d_%H-%M-%S")

    # Add date and time to filename
    filename = f'{task_name}_{dt_string}.docx'

    return send_file(buffer, as_attachment=True, attachment_filename=filename, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
